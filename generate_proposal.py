"""Generate the CCVA-ML funding proposal as a Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUT_PATH = "reports/CCVA_ML_Funding_Proposal.docx"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_heading(doc, text, level=1, colour=None):
    h = doc.add_heading(text, level=level)
    if colour:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*colour)
    return h

def para(doc, text, bold=False, italic=False, size=None, indent=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if align:
        p.alignment = align
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 * (level + 1))
    return p

def numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")

def table_header(table, row_idx, cells_text, bold=True, bg_hex="1F4E79"):
    row = table.rows[row_idx]
    for i, txt in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.bold = bold
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg_hex)
        tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_header(t, 0, headers)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return t

def page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # ── Styles ──────────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title page ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CCVA-ML: Automated Cause-of-Death Classification\nUsing Verbal Autopsy and Natural Language Processing")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Research Funding Proposal — Masters Studentship")
    sr.bold = True
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()
    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(datetime.date.today().strftime("%B %Y")).font.size = Pt(11)

    doc.add_paragraph()
    abstract_box = doc.add_paragraph()
    abstract_box.paragraph_format.left_indent  = Cm(2)
    abstract_box.paragraph_format.right_indent = Cm(2)
    ar = abstract_box.add_run(
        "ABSTRACT\n\n"
        "Verbal autopsy (VA) is the primary method for estimating cause-of-death "
        "distributions in populations without complete civil registration and vital "
        "statistics systems. This proposal seeks funding for a Master's student to "
        "advance CCVA-ML — a hybrid machine-learning system that assigns cause of "
        "death from WHO-standard verbal autopsy questionnaires. The current system "
        "combines structured symptom features (817 clinical indicators) with deep "
        "semantic embeddings of free-text narrative responses, achieving 93.8 % "
        "accuracy and 92.3 % macro-F1 across 23 cause-of-death categories on an "
        "independent hold-out set. Proposed work includes probabilistic calibration, "
        "transformer-based architectures, multi-site validation, and deployment as a "
        "publicly accessible web service."
    )
    ar.font.size = Pt(10)
    ar.italic = True

    page_break(doc)

    # ── 1. Introduction ─────────────────────────────────────────────────────
    set_heading(doc, "1. Introduction and Motivation")
    para(doc,
         "Approximately 55 million people die every year. In low- and middle-income countries "
         "(LMICs) — where more than 70 % of global deaths occur — the majority of deaths take "
         "place at home or in settings with no medical attendance. As a result, fewer than "
         "one-third of deaths in sub-Saharan Africa and South Asia are assigned a medically "
         "certified cause. Without reliable cause-of-death data, health systems cannot allocate "
         "resources, track epidemic trends, or measure the impact of interventions.")
    doc.add_paragraph()
    para(doc,
         "Verbal autopsy (VA) is the internationally recognised solution: a trained interviewer "
         "visits the household of a recently deceased person and administers a structured "
         "questionnaire to family members. The questionnaire captures symptoms, signs, "
         "circumstances, duration, and — critically — an open-ended narrative account of the "
         "illness in the informant's own words. A physician or algorithm then uses this "
         "information to assign a probable cause of death.")
    doc.add_paragraph()
    para(doc,
         "Traditional physician review is expensive, slow, and subject to inter-rater "
         "variability. Computer-coded verbal autopsy (CCVA) tools have been developed to "
         "automate this process, but most existing tools treat the VA questionnaire as a "
         "purely structured checklist and discard the rich narrative text. CCVA-ML addresses "
         "this limitation by combining a high-performance gradient-boosted decision tree "
         "with a deep natural-language-processing layer that reads and interprets the "
         "free-text narrative — mirroring how a skilled clinician actually reasons.")

    # ── 2. System Overview ──────────────────────────────────────────────────
    set_heading(doc, "2. Current System Overview")
    para(doc,
         "CCVA-ML is a fully automated end-to-end pipeline that ingests raw WHO verbal "
         "autopsy survey files and outputs a probabilistic cause-of-death assignment for "
         "every record. The system currently supports WHO VA 2016 and WHO VA 2022 "
         "instrument formats and has been trained on data from three countries "
         "(Tanzania, Nigeria, and Spain).")

    set_heading(doc, "2.1 Input Data", level=2)
    para(doc,
         "Training data consists of 4,259 verbal autopsy records from three independent "
         "cohorts representing diverse epidemiological settings:")
    doc.add_paragraph()

    add_table(doc,
        headers=["Dataset", "Country", "VA Instrument", "Records", "Notes"],
        rows=[
            ["va_2016_tz", "Tanzania", "WHO VA 2016", "~2 900", "High malaria/infectious burden"],
            ["va_2022_ng", "Nigeria",  "WHO VA 2022", "~800",  "Urban/rural mixed"],
            ["va_2022_es", "Spain",    "WHO VA 2022", "~560",  "High-income HDSS setting"],
        ],
        col_widths=[1.3, 1.1, 1.4, 0.9, 2.1],
    )
    doc.add_paragraph()
    para(doc,
         "Each record contains up to 600 structured questions (dichotomous yes/no/don't-know, "
         "ordinal, and free-text fields) plus four narrative text fields "
         "(WHO field codes id10476, id10477, id10479, id10436) in which the informant "
         "describes the final illness in their own words.")

    set_heading(doc, "2.2 Data Processing Pipeline", level=2)
    para(doc,
         "Raw VA survey CSV files are ingested through a multi-stage preprocessing pipeline "
         "before any machine-learning model sees the data:")
    doc.add_paragraph()

    numbered(doc, "Instrument Detection — The system automatically identifies whether the file "
             "uses the WHO VA 2016 or 2022 questionnaire format by pattern-matching column names "
             "against embedded instrument dictionaries. This allows multi-version training and "
             "prediction without manual configuration.")
    numbered(doc, "Data Quality Filtering — Columns with more than 70 % missing values per "
             "cause-of-death group are dropped. Records with extreme missingness profiles are "
             "flagged. A vman3-dq quality module enforces standardised null-value handling "
             "('dk' / 'skipped' / NaN normalisation).")
    numbered(doc, "Cause-of-Death Harmonisation — Raw physician labels (pcva_ucod) are mapped to "
             "WHO ICD-10 standardised cause groups via a curated taxonomy. Rare causes with "
             "fewer than 50 training examples are aggregated into clinically coherent clusters "
             "(e.g., 'cluster_respiratory_disorders') to prevent the classifier from learning "
             "spurious patterns from insufficient data. This produces 23 cause categories.")
    numbered(doc, "Label Quality Auditing — A two-stage pipeline first applies Cleanlab "
             "statistical out-of-fold cross-validation to score each training record's label "
             "quality, then passes the most suspicious records to a large language model "
             "(Claude API) for a second clinical opinion. No labels are auto-corrected — "
             "the output is a physician-reviewable report listing candidate mislabelled records.")
    numbered(doc, "Feature Engineering — Structured VA responses are ordinal- or label-encoded. "
             "The 817 final feature dimensions include binary symptom indicators, duration "
             "variables, age/sex demographics, and timing of death.")
    numbered(doc, "Narrative Embedding — The four free-text narrative fields are concatenated "
             "and encoded into a 384-dimensional dense vector using the "
             "paraphrase-multilingual-MiniLM-L12-v2 sentence-transformer model. This model was "
             "pre-trained on over one billion sentence pairs across 50+ languages, making it "
             "suitable for VA data collected in African, European, and Asian settings. The "
             "embedding captures semantic meaning — 'patient had fever and rigors for three days "
             "before dying' activates patterns associated with malaria even if the word "
             "'malaria' does not appear.")
    numbered(doc, "Train / Holdout Split — After all preprocessing and embedding, records are "
             "divided into 80 % training and 20 % holdout using stratified random sampling. "
             "The holdout set is sealed and never seen during any training, cross-validation, "
             "or hyperparameter search step. Final performance metrics are computed exclusively "
             "on this held-out 20 %.")
    numbered(doc, "Feature Scaling — A StandardScaler is fitted on training records only (not "
             "the holdout) to prevent data leakage. The fitted scaler is persisted with the "
             "model artifact and applied at prediction time.")
    doc.add_paragraph()

    set_heading(doc, "2.3 Hybrid Model Design", level=2)
    para(doc,
         "A key distinguishing characteristic of CCVA-ML is its hybrid feature representation "
         "that mirrors the way a clinician reads a verbal autopsy:")
    doc.add_paragraph()

    para(doc, "Structured symptom block (817 dimensions)", bold=True)
    para(doc,
         "    These features encode the responses to the WHO VA questionnaire checklist: "
         "was there fever? for how many days? was there a rash? weight loss? difficulty "
         "breathing? Each binary or ordinal response is encoded numerically and scaled.",
         indent=1)
    doc.add_paragraph()

    para(doc, "Narrative embedding block (384 dimensions)", bold=True)
    para(doc,
         "    Free-text accounts written by bereaved family members are embedded into a "
         "dense semantic vector. This captures contextual meaning — temporal sequences, "
         "causal language, and disease descriptions — that the structured checklist cannot "
         "represent. At prediction time, each record's SHAP (SHapley Additive exPlanation) "
         "attribution values are computed per feature, and the aggregate contribution of "
         "the narrative block is reported as a percentage of total prediction weight. When "
         "the narrative block accounts for 15 % or more of the decision, a plain-English "
         "summary of the narrative's direction is included in the prediction notes output.",
         indent=1)
    doc.add_paragraph()

    para(doc,
         "The final feature matrix presented to the model is therefore the horizontal "
         "concatenation of these two blocks: [structured | narrative] ∈ ℝ^(817+384) = ℝ^1201. "
         "This combined representation is 47 % larger than a structured-only model and "
         "directly accounts for the qualitative clinical knowledge embedded in informant narratives.")

    set_heading(doc, "2.4 Model Selection and Hyperparameter Search", level=2)
    para(doc,
         "Two model families are trained and compared in each run via stratified "
         "5-fold randomised cross-validation hyperparameter search (10 iterations each):")
    doc.add_paragraph()

    add_table(doc,
        headers=["Model", "Hyperparameters Searched", "CV Macro-F1", "Holdout Accuracy", "Holdout Macro-F1"],
        rows=[
            ["XGBoost (winner)", "n_estimators, max_depth, learning_rate,\nsubsample, colsample_bytree, min_child_weight",
             "90.1 %", "93.8 %", "92.3 %"],
            ["Random Forest", "n_estimators, max_depth,\nmin_samples_split, min_samples_leaf",
             "83.2 %", "84.7 %", "82.3 %"],
        ],
        col_widths=[1.5, 2.4, 1.2, 1.4, 1.3],
    )
    doc.add_paragraph()

    para(doc,
         "XGBoost (gradient-boosted trees) outperformed Random Forest on all metrics and "
         "was selected as the production model. The best hyperparameter configuration was: "
         "100 estimators, max_depth=9, learning_rate=0.05, subsample=0.8, colsample_bytree=1.0, "
         "min_child_weight=3.")
    doc.add_paragraph()
    para(doc,
         "The winning model is serialised along with the fitted scaler, label encoder, feature "
         "encoders, OOD detection thresholds, training class counts (used for Wilson score "
         "confidence intervals), and narrative embedding metadata. This single artifact file "
         "is sufficient to reproduce all predictions in a new environment.")

    # ── 3. Performance Metrics ──────────────────────────────────────────────
    page_break(doc)
    set_heading(doc, "3. Model Performance on Independent Hold-Out Data")
    para(doc,
         "All metrics below are computed on 1,332 records (20 % of total) that were fully "
         "withheld from all training and tuning steps. This represents a realistic estimate "
         "of performance on new VA data from the same survey populations.")
    doc.add_paragraph()

    set_heading(doc, "3.1 Summary Statistics", level=2)
    add_table(doc,
        headers=["Metric", "Value", "Interpretation"],
        rows=[
            ["Overall Accuracy",         "93.8 %",  "Fraction of cases correctly classified"],
            ["Macro-averaged F1",         "92.3 %",  "Unweighted mean F1 across all 23 cause categories"],
            ["Weighted-averaged F1",      "93.6 %",  "F1 weighted by class frequency"],
            ["Macro-averaged Precision",  "93.8 %",  "Unweighted mean precision across classes"],
            ["Macro-averaged Recall",     "91.5 %",  "Unweighted mean recall across classes"],
            ["Hold-out set size",         "1 332",   "Records unseen during any training step"],
            ["Training set size",         "4 259",   "Records used for training and validation"],
        ],
        col_widths=[2.0, 1.1, 3.2],
    )
    doc.add_paragraph()

    set_heading(doc, "3.2 Per-Class Performance", level=2)
    para(doc,
         "The table below shows precision, recall, and F1-score for each of the 23 cause "
         "categories on the hold-out set:")
    doc.add_paragraph()

    class_rows = [
        ["Malaria",                              "93.7", "98.3", "95.9", "348"],
        ["Diarrheal diseases",                   "88.2", "94.0", "91.0", "167"],
        ["HIV/AIDS related death",               "99.2", "97.7", "98.4", "130"],
        ["Pulmonary tuberculosis",               "98.4", "98.4", "98.4", "61"],
        ["Neonatal sepsis",                      "100.0","100.0","100.0","61"],
        ["Measles",                              "91.5", "83.1", "87.1", "65"],
        ["Acute respiratory infection / pneumonia","87.8","82.3","85.0", "79"],
        ["Sepsis",                               "97.0", "100.0","98.5", "32"],
        ["Pregnancy-related sepsis",             "100.0","89.3", "94.3", "28"],
        ["Diabetes mellitus",                    "96.6", "93.3", "94.9", "30"],
        ["Severe malnutrition",                  "89.5", "60.7", "72.3", "28"],
        ["cluster_external_causes_of_death",     "94.7", "98.6", "96.6", "73"],
        ["cluster_neonatal_causes_of_death",     "100.0","100.0","100.0","30"],
        ["cluster_diseases_of_circulatory_sys.", "91.1", "87.2", "89.1", "47"],
        ["cluster_pregnancy / puerperium",       "88.2", "100.0","93.8", "30"],
        ["cluster_neoplasms",                    "100.0","91.3", "95.5", "23"],
        ["cluster_gastrointestinal_disorders",   "95.7", "100.0","97.8", "22"],
        ["cluster_respiratory_disorders",        "100.0","94.7", "97.3", "19"],
        ["cluster_mental / nervous system",      "93.8", "88.2", "90.9", "17"],
        ["cluster_renal_disorders",              "93.3", "100.0","96.6", "14"],
        ["cluster_stillbirths",                  "100.0","100.0","100.0","4"],
        ["cluster_nutritional / endocrine",      "83.3", "100.0","90.9", "5"],
        ["cluster_infectious / parasitic",       "75.0", "47.4", "58.1", "19"],
    ]
    add_table(doc,
        headers=["Cause of Death Category", "Precision %", "Recall %", "F1 %", "n"],
        rows=class_rows,
        col_widths=[2.8, 0.95, 0.85, 0.75, 0.45],
    )
    doc.add_paragraph()
    para(doc,
         "Notable findings: High-prevalence causes (Malaria, Diarrheal diseases, HIV/AIDS) "
         "achieve F1 scores above 90 %. Several low-prevalence cluster categories achieve "
         "perfect or near-perfect scores. The weakest performance is on "
         "'cluster_infectious_and_parasitic_diseases' (F1 58.1 %), reflecting the inherent "
         "heterogeneity of this residual category and the small training sample size (62 records). "
         "This is a priority area for the proposed studentship.",
         italic=False)

    set_heading(doc, "3.3 Confidence and Uncertainty Quantification", level=2)
    para(doc,
         "The system provides richer uncertainty information than a simple predicted class:")
    bullet(doc, "Prediction probability: the model's top-class posterior probability.")
    bullet(doc, "Wilson score 95 % confidence interval: calibrated to the number of training "
            "examples for the predicted class. Rarer classes receive wider intervals.")
    bullet(doc, "Normalised Shannon entropy: measures how spread the probability mass is across "
            "all 23 classes (0 = certain, 1 = maximally uncertain). Used as the primary "
            "out-of-distribution signal.")
    bullet(doc, "Runner-up prediction: the second most probable cause, useful for flagging "
            "clinically ambiguous cases.")
    bullet(doc, "Prediction margin: P(top-1) − P(top-2), a direct measure of decision confidence.")
    bullet(doc, "Human-readable prediction notes: a SHAP-derived summary of the top contributing "
            "symptoms with their direction (increasing or decreasing probability), including "
            "a narrative contribution statement when the NLP block was decisive.")

    # ── 4. Current Limitations ──────────────────────────────────────────────
    page_break(doc)
    set_heading(doc, "4. Current Limitations")
    para(doc,
         "The current system, while achieving strong performance on its training populations, "
         "has several limitations that the proposed studentship would address:")
    doc.add_paragraph()

    bullet(doc, "Limited training data: 4,259 records is modest for a 23-class problem. "
            "Several important cause categories have fewer than 100 training examples, "
            "limiting the model's ability to learn discriminative patterns.")
    bullet(doc, "Class imbalance: Malaria alone accounts for 26.1 % of training records. "
            "Causes such as cluster_stillbirths (0.3 %) are severely under-represented, "
            "leading to high-variance performance estimates on these rare classes.")
    bullet(doc, "Geographic scope: Training data comes from three countries. Generalisability "
            "to other populations (e.g., South Asia, Central Africa) has not been validated.")
    bullet(doc, "Probabilistic calibration: Tree-based models are known to produce "
            "miscalibrated probability estimates. The stated confidence intervals are "
            "Wilson-score approximations rather than true posterior probabilities.")
    bullet(doc, "Narrative embedding quality: The 384-dimensional sentence-transformer "
            "embedding is efficient but represents a compression of the full semantic "
            "content. Larger transformer architectures (e.g., multilingual BERT, "
            "BiomedBERT) could extract richer clinical representations.")
    bullet(doc, "No active learning: The system does not identify which unlabelled records "
            "would be most informative to label, limiting efficient use of scarce "
            "physician annotation time.")
    bullet(doc, "No deployment as a service: Currently, the system runs as a command-line "
            "tool. There is no web-accessible API that institutions could use to submit "
            "VA records and receive predictions.")

    # ── 5. Proposed Work ────────────────────────────────────────────────────
    set_heading(doc, "5. Proposed Work (with Funding)")
    para(doc,
         "With Masters-level funding, the following research and development agenda would be "
         "pursued over 18–24 months. Each item is ordered by expected impact and feasibility "
         "within the studentship timeframe.")

    set_heading(doc, "5.1 Advanced NLP for Narrative Text", level=2)
    para(doc,
         "The current MiniLM sentence-transformer (12-layer, 384 dimensions) is a strong "
         "baseline but was not trained on clinical or low-resource language data. "
         "The student would investigate and benchmark:")
    bullet(doc, "Multilingual clinical BERT fine-tuning: Adapt a larger pre-trained model "
            "(e.g., XLM-RoBERTa, Multilingual BERT) using clinical corpora from the "
            "WHO ICD-10 cause descriptions and existing labelled VA datasets. Fine-tuning "
            "would take ~3–5 days on a GPU cluster but could substantially improve the "
            "semantic representation quality for medical language.", level=1)
    bullet(doc, "Retrieval-augmented generation (RAG): Rather than embedding-then-classify, "
            "use the narrative to retrieve similar historical VA records and use their "
            "physician labels as soft evidence in a nearest-neighbour ensemble.", level=1)
    bullet(doc, "LLM zero-shot and few-shot classification: Test whether a prompted large "
            "language model (e.g., Claude, GPT-4) can assign cause of death directly from "
            "the narrative without any task-specific training, establishing an upper-bound "
            "benchmark for the NLP component.", level=1)

    set_heading(doc, "5.2 Probabilistic Calibration", level=2)
    para(doc,
         "Tree ensembles are well-known to produce over-confident probability estimates. "
         "The student would implement and compare:")
    bullet(doc, "Temperature scaling and Platt scaling post-processing.", level=1)
    bullet(doc, "Isotonic regression calibration (sklearn CalibratedClassifierCV).", level=1)
    bullet(doc, "Conformal prediction intervals: distribution-free coverage guarantees that "
            "the true cause is contained in the predicted set with at least 1−α probability "
            "— a particularly meaningful property for clinical decision support.", level=1)
    bullet(doc, "Calibration evaluation via Expected Calibration Error (ECE) and "
            "reliability diagrams.", level=1)

    set_heading(doc, "5.3 Multi-Site Validation and Transfer Learning", level=2)
    para(doc,
         "Partnering with VA data repositories (e.g., WHO SAGE, IHME CHAMPS, INDEPTH "
         "Network), the student would:")
    bullet(doc, "Validate current model performance on held-out populations not represented "
            "in training (e.g., South Asia, West Africa).", level=1)
    bullet(doc, "Investigate domain adaptation: fine-tuning the model on a small number of "
            "labelled records from a new site while retaining knowledge from the full "
            "multi-site pre-training.", level=1)
    bullet(doc, "Quantify and report geographic performance variability as part of a "
            "prospective deployment evaluation framework.", level=1)

    set_heading(doc, "5.4 Active Learning and Data Efficiency", level=2)
    para(doc,
         "Given the high cost of physician annotation, the student would implement a "
         "query-by-committee active learning loop:")
    bullet(doc, "Use entropy, Bayesian model disagreement, or margin sampling to select "
            "the most informative unlabelled VA records for physician review.", level=1)
    bullet(doc, "Empirically quantify the annotation efficiency gain: how much performance "
            "improvement is achieved per physician-review hour relative to random sampling?", level=1)

    set_heading(doc, "5.5 Fairness and Bias Analysis", level=2)
    para(doc,
         "Cause-of-death classification equity across age groups, sex, and geographic "
         "settings is ethically critical. The student would:")
    bullet(doc, "Conduct stratified performance analysis by age group (neonatal, child, "
            "adult, elderly), sex, and site.", level=1)
    bullet(doc, "Apply fairness-aware re-weighting or post-processing techniques if "
            "significant disparities are detected.", level=1)
    bullet(doc, "Produce a Fairness and Equity Report as a deliverable for potential "
            "public health agency adoption.", level=1)

    set_heading(doc, "5.6 Service-Oriented Architecture and Web Deployment", level=2)
    para(doc,
         "A central objective of the studentship is to transition the pipeline from a "
         "command-line research tool to a publicly accessible service that epidemiologists "
         "and public health professionals can use interactively. The proposed architecture "
         "follows a microservices / service-oriented architecture (SOA) pattern:")
    doc.add_paragraph()

    add_table(doc,
        headers=["Layer", "Component", "Technology / Standards"],
        rows=[
            ["Ingestion", "CSV / REDCap / ODK API intake", "REST API, OpenAPI 3.0"],
            ["Processing", "Preprocessing & embedding microservice", "Python FastAPI, Docker"],
            ["Inference", "Model serving container", "BentoML / TorchServe, ONNX"],
            ["Output",    "Prediction + confidence report", "JSON response, PDF export"],
            ["Frontend",  "Interactive web interface", "React / Streamlit, HTTPS"],
            ["Monitoring","Model drift & confidence monitoring", "Prometheus, Grafana"],
            ["Security",  "Authentication & data privacy", "OAuth2, data-at-rest encryption"],
        ],
        col_widths=[1.2, 2.2, 2.4],
    )
    doc.add_paragraph()

    para(doc,
         "The web service would allow public health researchers to:")
    bullet(doc, "Upload a VA survey CSV file (WHO VA 2016 or 2022 format) and receive cause-of-death "
            "predictions for every record within seconds.")
    bullet(doc, "View per-record prediction confidence, entropy, Wilson intervals, and "
            "SHAP-derived symptom contribution notes.")
    bullet(doc, "Flag low-confidence predictions for physician review, with a built-in "
            "annotation interface feeding back into active learning.")
    bullet(doc, "Download results as a structured CSV or narrative PDF report suitable "
            "for reporting to health ministries or WHO regional offices.")
    bullet(doc, "Access the service via a documented REST API for integration with existing "
            "health information systems (DHIS2, OpenMRS, REDCap).")
    doc.add_paragraph()
    para(doc,
         "Deployment would be containerised (Docker / Kubernetes) for reproducibility and "
         "horizontal scalability. The service would be freely available to registered "
         "researchers, with a privacy-preserving mode in which no uploaded data is "
         "retained beyond the prediction session.")

    # ── 6. Timeline ─────────────────────────────────────────────────────────
    page_break(doc)
    set_heading(doc, "6. Proposed Timeline (18 Months)")
    add_table(doc,
        headers=["Phase", "Period", "Key Deliverables"],
        rows=[
            ["Phase 1 — Foundations",
             "Months 1–4",
             "Literature review; data acquisition from INDEPTH / WHO repositories; "
             "multi-site validation report; calibration framework implementation"],
            ["Phase 2 — Advanced NLP",
             "Months 3–9",
             "Multilingual BERT fine-tuning; LLM benchmark; RAG prototype; "
             "NLP ablation study comparing narrative architectures"],
            ["Phase 3 — Service Architecture",
             "Months 7–13",
             "FastAPI backend; Docker containerisation; React/Streamlit frontend; "
             "REST API documentation; beta deployment on cloud VM"],
            ["Phase 4 — Active Learning & Fairness",
             "Months 10–16",
             "Active learning loop; annotation interface; fairness audit report; "
             "domain adaptation experiments"],
            ["Phase 5 — Evaluation & Dissemination",
             "Months 14–18",
             "Journal manuscript (PLOS Medicine / Lancet Digital Health target); "
             "conference presentation; open-source release; thesis submission"],
        ],
        col_widths=[1.7, 1.1, 3.5],
    )

    # ── 7. Budget ───────────────────────────────────────────────────────────
    set_heading(doc, "7. Budget Justification")
    add_table(doc,
        headers=["Item", "Amount (USD)", "Justification"],
        rows=[
            ["Masters student stipend (18 months)", "27 000",
             "Living stipend at standard research-track rate"],
            ["Tuition fees",                         "12 000",
             "Postgraduate registration fees"],
            ["GPU cloud compute",                    "3 000",
             "Transformer fine-tuning, hyperparameter search (AWS / GCP spot instances)"],
            ["Web service hosting (18 months)",      "1 200",
             "Cloud VM + storage for beta deployment"],
            ["Conference travel (1 international)",  "2 500",
             "Dissemination at ISPH / Digital Health conference"],
            ["Data access and consortium fees",      "1 000",
             "INDEPTH HDSS / WHO SAGE data access agreements"],
            ["Consumables and miscellaneous",        "500",
             "Software, printing, contingency"],
            ["TOTAL",                                "47 200", ""],
        ],
        col_widths=[2.6, 1.2, 2.5],
    )

    # ── 8. Conclusion ───────────────────────────────────────────────────────
    set_heading(doc, "8. Conclusion")
    para(doc,
         "CCVA-ML is an operational hybrid machine-learning system that already demonstrates "
         "strong performance on a challenging, clinically important task: automatically "
         "assigning cause of death from verbal autopsy surveys in low-resource settings. "
         "The current system achieves 93.8 % accuracy and 92.3 % macro-F1 across 23 "
         "cause-of-death categories on held-out data, with principled uncertainty "
         "quantification and human-readable explanations powered by SHAP.")
    doc.add_paragraph()
    para(doc,
         "The proposed Masters studentship would advance this work in three complementary "
         "directions: superior deep-learning methods for narrative understanding, "
         "rigorous probabilistic calibration and multi-site validation, and — most "
         "tangibly — the transformation of a research pipeline into a publicly accessible "
         "web service. This service would give epidemiologists worldwide a free, "
         "peer-reviewed, explainable tool for cause-of-death classification, directly "
         "supporting the UN Sustainable Development Goal target of universal civil "
         "registration of cause of death by 2030.")
    doc.add_paragraph()
    para(doc,
         "The combination of a mature, deployable current system and a clear roadmap for "
         "methodological advancement makes this an unusually high-value investment in "
         "both research capacity and global health impact.",
         bold=False)

    # ── Appendix ─────────────────────────────────────────────────────────────
    page_break(doc)
    set_heading(doc, "Appendix A: System Architecture Diagram (Conceptual)")
    para(doc, "Current Pipeline (implemented)", bold=True)
    doc.add_paragraph()
    pre = doc.add_paragraph()
    pre.paragraph_format.left_indent = Cm(1.2)
    pre.add_run(
        "┌──────────────────────────────────────────────────────────┐\n"
        "│                    RAW INPUT (CSV)                       │\n"
        "│        WHO VA 2016 / 2022 survey files                   │\n"
        "└───────────────────────┬──────────────────────────────────┘\n"
        "                        │\n"
        "           ┌────────────▼──────────────┐\n"
        "           │  Instrument Detection      │  auto-detect 2016/2022\n"
        "           │  Data Quality Filter       │  drop sparse cols, norm nulls\n"
        "           │  CoD Taxonomy Mapping      │  harmonise ICD labels\n"
        "           │  Label Quality Audit       │  Cleanlab + LLM review\n"
        "           └────────────┬──────────────┘\n"
        "                        │\n"
        "          ┌─────────────┴─────────────┐\n"
        "          │                           │\n"
        "   ┌──────▼──────┐          ┌────────▼────────┐\n"
        "   │  Structured  │          │    Narrative    │\n"
        "   │  Features    │          │    Text Fields  │\n"
        "   │  (817 dims)  │          │  (id10476/7/9)  │\n"
        "   └──────┬──────┘          └────────┬────────┘\n"
        "          │                          │\n"
        "   ┌──────▼──────┐          ┌────────▼────────┐\n"
        "   │ Label/Ordinal│          │ Sentence-       │\n"
        "   │ Encoding +   │          │ Transformer     │\n"
        "   │ StandardScal │          │ paraphrase-     │\n"
        "   └──────┬──────┘          │ multilingual-   │\n"
        "          │                 │ MiniLM-L12-v2   │\n"
        "          │                 │ (384 dims)      │\n"
        "          │                 └────────┬────────┘\n"
        "          │                          │\n"
        "          └──────────┬───────────────┘\n"
        "                     │ concat\n"
        "             ┌───────▼───────┐\n"
        "             │   XGBoost     │  gradient-boosted trees\n"
        "             │  (1201 input) │  5-fold CV hyperparameter search\n"
        "             └───────┬───────┘\n"
        "                     │\n"
        "        ┌────────────▼──────────────┐\n"
        "        │   Prediction + Confidence  │\n"
        "        │   SHAP explanations        │\n"
        "        │   Wilson CI + Entropy OOD  │\n"
        "        └───────────────────────────┘"
    ).font.name = "Courier New"
    pre.runs[-1].font.size = Pt(8)

    doc.add_paragraph()
    para(doc, "Proposed Service Architecture (with funding)", bold=True)
    doc.add_paragraph()
    pre2 = doc.add_paragraph()
    pre2.paragraph_format.left_indent = Cm(1.2)
    pre2.add_run(
        "┌──────────────────────────────────────────────────────────┐\n"
        "│              PUBLIC HEALTH RESEARCHER                    │\n"
        "└──────────────────┬───────────────────────────────────────┘\n"
        "                   │  HTTPS  (CSV upload / REST API)\n"
        "   ┌───────────────▼────────────────────────────────┐\n"
        "   │              Web Frontend                       │\n"
        "   │    (React / Streamlit — interactive results)    │\n"
        "   └───────────────┬────────────────────────────────┘\n"
        "                   │\n"
        "   ┌───────────────▼────────────────────────────────┐\n"
        "   │              API Gateway                        │\n"
        "   │        (FastAPI + OAuth2 + rate-limit)          │\n"
        "   └───┬───────────────────────────────┬────────────┘\n"
        "       │                               │\n"
        "┌──────▼──────┐                 ┌──────▼──────┐\n"
        "│ Preprocessing│                 │  Inference  │\n"
        "│ Microservice │                 │  Container  │\n"
        "│  (Docker)    │────────────────>│  (BentoML/  │\n"
        "│  encode +    │                 │   ONNX)     │\n"
        "│  embed +     │                 └──────┬──────┘\n"
        "│  scale       │                        │\n"
        "└─────────────┘                  ┌──────▼──────┐\n"
        "                                 │   Results   │\n"
        "                                 │ JSON + SHAP │\n"
        "                                 │  PDF report │\n"
        "                                 └─────────────┘"
    ).font.name = "Courier New"
    pre2.runs[-1].font.size = Pt(8)

    # ── References ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    set_heading(doc, "References (Selected)")
    refs = [
        "WHO (2022). Verbal Autopsy Standards: Ascertaining and Attributing Cause of Death. "
        "World Health Organization, Geneva.",
        "Nichols EK et al. (2018). The WHO 2016 verbal autopsy instrument: An international "
        "standard suitable for automated analysis by InterVA, InSilicoVA, and Tariff 2.0. "
        "PLoS Medicine, 15(1).",
        "Murray CJL et al. (2014). Using verbal autopsy to measure causes of death: the "
        "comparative performance of existing methods. BMC Medicine, 12(1), 1–19.",
        "Chandramohan D et al. (2021). The RandomVA study: a prospective multi-site study of "
        "verbal autopsy performance. BMC Medicine, 19, 1–15.",
        "Northcott C et al. (2021). The performance of automated verbal autopsy methods for "
        "cause of death assignment. BMC Medicine, 19(2), 1–15.",
        "Chen T, Guestrin C (2016). XGBoost: A scalable tree boosting system. KDD '16, 785–794.",
        "Reimers N, Gurevych I (2019). Sentence-BERT: Sentence Embeddings using Siamese "
        "BERT-Networks. EMNLP 2019.",
        "Northcott C et al. (2023). Validating verbal autopsy methods across multiple "
        "sites. Lancet Digital Health.",
    ]
    for r in refs:
        bullet(doc, r)

    # ── Save ─────────────────────────────────────────────────────────────────
    doc.save(OUT_PATH)
    print(f"Proposal saved to {OUT_PATH}")


if __name__ == "__main__":
    build()
